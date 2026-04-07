#!/usr/bin/env python3
from docx import Document
import chromadb
import requests

def get_embeddings(prompts, model="herald/dmeta-embedding-zh:latest"):
    results = []
    for prompt in prompts:
        r = requests.post(
            "http://localhost:11434/api/embeddings",
            json={"model": model, "prompt": prompt},
            timeout=60
        )
        results.append(r.json()["embedding"])
    return results

def chunk_text(text, size=600, overlap=100):
    chunks = []
    start = 0
    L = len(text)
    while start < L:
        end = start + size
        if end >= L:
            chunks.append(text[start:])
            break
        chunk = text[start:end]
        bp = chunk.rfind('\n')
        if bp < size * 0.3:
            for p in "。；！？":
                b = chunk.rfind(p)
                if b > size * 0.5:
                    bp = b
                    break
        actual = start + bp + 1 if bp > size * 0.5 else end
        chunks.append(text[start:actual])
        start = actual - overlap
    return chunks

def extract_docx(fname):
    doc = Document(fname)
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [c.text.strip().replace('\n', ' ') for c in row.cells if c.text.strip()]
            if cells:
                rows.append(' | '.join(cells))
        if rows:
            tables.append('\n'.join(rows))
    return paras, tables

chunks_data = []

files = [
    ('/Users/momobalap/.openclaw/workspace/cvd-kg/Array1_CVD设备管控作业指导书.docx', 'CVD设备管控作业指导书', 'CVD'),
    ('/Users/momobalap/.openclaw/workspace/cvd-kg/Array1_薄膜产品异常处理作业指导书.docx', '薄膜产品异常处理作业指导书', '薄膜异常'),
]

for fname, src, sec in files:
    print(f'Extracting {src}...')
    paras, tables = extract_docx(fname)
    print(f'  {len(paras)} paragraphs, {len(tables)} tables')
    
    combined = '\n'.join(paras)
    for c in chunk_text(combined, 600, 100):
        if len(c) > 50:
            chunks_data.append((c, src, sec, 'para'))
    
    for t in tables:
        if len(t) > 30:
            for c in chunk_text(t, 600, 100):
                if len(c) > 30:
                    chunks_data.append((c, src, sec, 'table'))

print(f'Total chunks: {len(chunks_data)}')

client = chromadb.PersistentClient(path='/Users/momobalap/.openclaw/workspace/cvd-kg/chroma_sop_v3')
try:
    client.delete_collection('sop_full')
except:
    pass

col = client.create_collection(name='sop_full', metadata={'description': 'SOP完整文档 含表格 600/100 v3'})

batch_size = 10
total = len(chunks_data)
for i in range(0, total, batch_size):
    batch = chunks_data[i:i+batch_size]
    texts = [c[0] for c in batch]
    ids = [f'c_{i+j}_{batch[j][3]}' for j in range(len(batch))]
    metas = [{'source': c[1], 'section': c[2], 'type': c[3]} for c in batch]
    print(f'Batch {i//batch_size+1}/{(total+batch_size-1)//batch_size}...', end=' ', flush=True)
    embs = get_embeddings(texts)
    col.add(ids=ids, embeddings=embs, documents=texts, metadatas=metas)
    print(f'+{len(batch)}')

print(f'Stored: {col.count()}')
